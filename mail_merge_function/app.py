import pandas as pd
import boto3
import os
from io import StringIO
import logging
from docx import Document

"""
The lambda_handler function reads a CSV file from an S3 bucket, performs a mail merge with a Word
document template, and generates a merged Word document.

:param event: The `event` parameter is used to pass data to the Lambda function when it is
triggered. It can contain information such as the event type, event source, and any additional data
relevant to the event
:param context: The `context` parameter is a context object provided by AWS Lambda. It contains
information about the runtime environment and function invocation. It can be used to access the AWS
request ID, function name, memory limit, and other useful information
"""

def lambda_handler(event, context):
    

#   Find object in our s3 buckets. 
    s3 = boto3.client('s3')
    logger = logging.getLogger()
    logger.setLevel(logging.INFO)
    
    # Label all of our environments that we initially use
    task_route = os.environ.get('task_route')
    bucket_name = os.environ.get('BUCKET_NAME')
    dataframe_file_s3 = os.environ.get('csv_file')
    template = os.environ.get('template')
    output = os.environ.get('output')
    
    # Log the information to be seen more transparently
    logger.info(f"Reading Bucket: {bucket_name}")
    logger.info(f"file: {dataframe_file_s3}")
    logger.info(f"file: {template}")
    
    print(f"Reading file: {dataframe_file_s3} to be merged with {template} resulting in the file {output}")
    
    # Get data file that holds the needed information
    df_object = s3.get_object(Bucket = bucket_name, Key = dataframe_file_s3)
    df_body = df_object['Body'].read().decode('utf-8')
    
    df_io = StringIO(df_body)
    df = pd.read_csv(df_io)
    
    df.to_csv('/tmp/dataset.csv', index=False)
    
    data = pd.read_csv('/tmp/dataset.csv')

        
    data['PHONE'] = data['PHONE'].astype(str).str.replace('-', '')
    data['PHONE'].fillna('', inplace=True)
    
    # Iterate through each row in the CSV
    print(data.head(10))
    
    print('Confirm Dataset is read, now we go to download.')
    
    # Load your Word document template

    merged_document = Document()
    print('New Document Created.')
    
    # Load your Word document template
    # Iterate through each row in the CSV
    for index, row in data.iterrows():
        document = Document(f'{task_route}{template}')
        # Replace placeholders with values from the current row
        for key, value in row.items():
            placeholder = f'«{key}»'
            for para in document.paragraphs:
                for run in para.runs:
                    if placeholder in run.text:
                        if key=='USER_PIN':
                            run.text = run.text.replace(placeholder, str(value).zfill(4))
                        elif key=='PHONE':
                            run.text = run.text.replace(placeholder, str(value).zfill(0))
                        else:
                            run.text = run.text.replace(placeholder, str(value))
                            
    


        # Append the merged document to the final document
        for element in document.element.body:
            merged_document.element.body.append(element)
    
    print('Document Merged.')  
    # Remove section breaks from the merged document
    # section_breaks = merged_document.paragraphs
    # for para in section_breaks:
    #     if para.text == '\x0c':
    #         para.clear()
            
    # for para in merged_document.paragraphs:
    #     for run in para.runs:
    #         run.text = run.text.replace('|anan', '|')           
    # merged_document.save(output)
    
    # with open(output, 'w') as f:
    #     for paragraph in merged_document.paragraphs:
    #         f.write(paragraph.text)
    #     logger.info(f"Converting to txt: {f}")
        
    #   Augmented Code Here
    
    merged_document.save(output)
    with open(output, 'w') as f:        
        section_breaks = merged_document.paragraphs
        for para in section_breaks:
            text = para.text.replace('|anan', '|')
            text = text.replace('.0', '')
            f.write(text + '\n')     
                             
            if para.text == '\x0c':
                para.clear()  
    
    #   Augmented Code Finished here
        
    
    
    deposited_name = os.environ.get('deposited_name') 
    s3 = boto3.client('s3')
    with open(output, 'rb') as data_merged:
        s3.upload_fileobj(data_merged, bucket_name, deposited_name)
    logger.info(f"Uploading file: {deposited_name}")
    return {
    'statusCode': 200
    }